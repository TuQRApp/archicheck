from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Márgenes ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# --- Paleta ---
NEGRO      = RGBColor(0x1A, 0x1A, 0x1A)
AZUL       = RGBColor(0x1E, 0x40, 0xAF)   # azul ArchiCheck
GRIS       = RGBColor(0x6B, 0x72, 0x80)
GRIS_CLARO = RGBColor(0xF3, 0xF4, 0xF6)
ROJO       = RGBColor(0xDC, 0x26, 0x26)

# ─────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────

def set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1, color=AZUL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size = 16 if level == 1 else 13
    set_run_font(run, size=size, bold=True, color=color)
    return p

def body(text, indent=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_run_font(run, size=11, color=NEGRO)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_run_font(r1, size=11, bold=True, color=NEGRO)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=NEGRO)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=NEGRO)
    return p

def tag_line(label, texto):
    """Línea con etiqueta en azul y texto normal."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10, bold=True, color=AZUL)
    r2 = p.add_run(texto)
    set_run_font(r2, size=11, color=NEGRO)
    return p

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 72)
    set_run_font(run, size=8, color=GRIS)
    return p

def nota(texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto)
    set_run_font(run, size=10, color=GRIS, italic=True)
    return p

# ─────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run("ArchiCheck")
set_run_font(r, size=26, bold=True, color=AZUL)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run("Plan de trabajo · Startup Chile Ignite")
set_run_font(r2, size=14, color=GRIS)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(16)
r3 = p3.add_run("4 meses · CLP $30 millones · Agosto–Diciembre 2026")
set_run_font(r3, size=11, color=GRIS, italic=True)

separator()

# Nota introductoria
body(
    "El programa Ignite entrega financiamiento equity-free, mentorías especializadas, "
    "conexión con corporativos e inversionistas, y actividades de aceleración intensiva. "
    "Este plan estructura el trabajo en 5 tareas principales que absorben esas actividades "
    "del programa — mentores, redes y plan de negocios no son actividades aparte, sino "
    "componentes explícitos dentro de cada tarea.",
    space_after=12
)

# ─────────────────────────────────────────
# TABLA DE DEPENDENCIAS (resumen visual)
# ─────────────────────────────────────────

heading("Estructura y dependencias", level=1)

body(
    "T1 y T2 arrancan el día 1 en paralelo. T3 depende de precisión suficiente en T1 + "
    "señales de T2. T4 puede arrancar con el demo actual antes de que T1 esté completa. "
    "T5 consolida las cuatro anteriores. Las actividades de mentores y redes corren "
    "transversal a todo el programa — no son eventos puntuales.",
    space_after=8
)

# Tabla resumen
table = doc.add_table(rows=6, cols=5)
table.style = "Table Grid"
headers = ["Tarea", "Meses", "Paralela con", "Depende de", "Objetivo clave"]
widths   = [Cm(3.5), Cm(2.2), Cm(3.0), Cm(2.8), Cm(5.5)]

for i, (h, w) in enumerate(zip(headers, widths)):
    cell = table.rows[0].cells[i]
    cell.width = w
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_run_font(run, size=10, bold=True, color=AZUL)
    # fondo azul oscuro para encabezado
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DBEAFE")
    tcPr.append(shd)

rows_data = [
    ("T1 · Extracción", "1–2", "T2", "—", "Precisión geométrica sobre planos chilenos"),
    ("T2 · Validación", "1–3", "T1, T4", "—", "5–8 arquitectos en piloto, feedback documentado"),
    ("T3 · Informe",    "2–3", "—",    "T1 + T2", "Ciclo completo: plano → informe DOM"),
    ("T4 · Comercial",  "2–4", "T1, T2", "—",    "Plan de negocios + primeras ventas reales"),
    ("T5 · Cierre",     "4",   "—",    "T1–T4",  "Métricas consolidadas + pipeline inversión"),
]

for r_idx, row_data in enumerate(rows_data, start=1):
    row = table.rows[r_idx]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_run_font(run, size=10, color=NEGRO)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ─────────────────────────────────────────
# TAREAS
# ─────────────────────────────────────────

tareas = [
    {
        "titulo": "T1 — Motor de extracción confiable",
        "periodo": "Meses 1–2",
        "desc": (
            "Integrar y calibrar Floor Plan API (y/o alternativas identificadas en el benchmark) "
            "sobre 10+ planos reales chilenos de distintos tipos (oficinas, comercio, vivienda). "
            "Definir métricas internas de precisión: IoU de muros y F1 de vanos, medidas contra "
            "ground truth propio. Este es el cuello de botella tecnológico que habilita T3."
        ),
        "mentores": (
            "Conectar con al menos 1 mentor técnico (IA / visión computacional) para auditar "
            "la arquitectura de extracción y la elección de herramientas. Priorizar mentores "
            "del ecosistema Startup Chile con experiencia en computer vision o construcción tech."
        ),
        "entregables": [
            "Dataset ground truth: 10+ planos anotados con geometría verificada.",
            "Informe de precisión técnica: IoU y F1 por tipo de plano y herramienta.",
            "Decisión documentada de stack (Floor Plan API / alternativa) con justificación.",
        ],
    },
    {
        "titulo": "T2 — Validación con arquitectos",
        "periodo": "Meses 1–3 (paralela a T1)",
        "desc": (
            "Onboardear 5–8 arquitectos en piloto (pagado o simbólico). El foco es capturar "
            "qué hallazgos validan como útiles, cuáles rechazan y por qué — insumo directo "
            "para calibrar T1 y definir el formato de T3."
        ),
        "mentores": (
            "Usar la red de Startup Chile (founders, corporativos, actividades del programa) "
            "para acceso a estudios de arquitectura y contactos en DOMs. Al menos 1 mentor "
            "con experiencia en ventas B2B o en el sector construcción/inmobiliario."
        ),
        "entregables": [
            "8+ entrevistas documentadas (problema, flujo actual, disposición a pagar).",
            "Matriz de feedback por tipo de observación (acepta / rechaza / modifica).",
            "3+ pilotos activos con acuerdo firmado (aunque sea piloto gratuito).",
        ],
    },
    {
        "titulo": "T3 — Informe de cumplimiento exportable",
        "periodo": "Meses 2–3",
        "desc": (
            "Cerrar el ciclo completo del producto: plano → extracción geométrica → revisión "
            "OGUC/PRC → informe Word/PDF con citas normativas exactas por artículo, listo para "
            "adjuntar a un expediente DOM. Es la entrega de valor que convierte el prototipo "
            "en algo que el arquitecto puede cobrarle a su cliente."
        ),
        "mentores": (
            "1 mentor de producto/UX para validar el formato del informe con el usuario real "
            "antes de construirlo. Si el programa conecta con corporativos del sector "
            "inmobiliario, usar esas reuniones para testear el informe como artefacto."
        ),
        "entregables": [
            "Template de informe validado por al menos 3 arquitectos.",
            "Documentación de las reglas normativas implementadas (artículos OGUC/PRC codificados).",
            "Integración del informe en el flujo del producto (generación automática desde la plataforma).",
        ],
    },
    {
        "titulo": "T4 — Comercialización y plan de negocios",
        "periodo": "Meses 2–4 (paralela)",
        "desc": (
            "Definir precio, canal y propuesta de valor concreta (qué le ahorra en tiempo y "
            "dinero al arquitecto, medido con datos reales de T2). Ejecutar primeras ventas. "
            "Desarrollar el modelo financiero del negocio usando los talleres y mentorías del programa."
        ),
        "mentores": (
            "1 mentor de go-to-market o ventas SaaS B2B. Usar las actividades de aceleración "
            "de Startup Chile (talleres de pricing, pitch, modelo de negocio) como insumo directo "
            "para construir el plan. Networking con corporativos del ecosistema como potenciales "
            "clientes o aliados de distribución."
        ),
        "entregables": [
            "Plan de negocios: modelo de ingresos, proyección a 12 meses, CAC/LTV estimado.",
            "3–5 ventas documentadas (contrato o factura).",
            "MRR demostrable al cierre del programa.",
            "Benchmark competitivo actualizado (Revi, revisores manuales, herramientas internacionales).",
        ],
    },
    {
        "titulo": "T5 — Cierre e inicio de levantamiento de capital",
        "periodo": "Mes 4",
        "desc": (
            "Consolidar métricas del programa (precisión técnica, usuarios activos, MRR) y abrir "
            "conversaciones con inversionistas usando los datos reales como tracción. No se postula "
            "a Growth ahora — el objetivo es tener todo listo para levantar una ronda seed una vez "
            "cerrado el Ignite."
        ),
        "mentores": (
            "Usar las conexiones de Startup Chile con ángeles e inversionistas early-stage para "
            "primeras reuniones formales. Al menos 1 mentor con experiencia en levantamiento de "
            "capital seed en LatAm. Objetivo: 2–3 reuniones con inversionistas antes del Demo Day."
        ),
        "entregables": [
            "Pitch deck actualizado con datos reales del programa.",
            "Informe de cierre Ignite: métricas vs. metas definidas en T1–T4.",
            "Pipeline de inversión documentado: contactos, estado de conversación, próximos pasos.",
        ],
    },
]

for tarea in tareas:
    separator()
    heading(tarea["titulo"], level=1)
    tag_line("Período:", tarea["periodo"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after  = Pt(8)
    p_desc.paragraph_format.space_before = Pt(0)
    run = p_desc.add_run(tarea["desc"])
    set_run_font(run, size=11, color=NEGRO)

    # Mentores y redes
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_after  = Pt(3)
    p_m.paragraph_format.space_before = Pt(0)
    r_label = p_m.add_run("Mentores y redes (Startup Chile)  ")
    set_run_font(r_label, size=11, bold=True, color=AZUL)

    p_m2 = doc.add_paragraph()
    p_m2.paragraph_format.left_indent  = Cm(0.6)
    p_m2.paragraph_format.space_after  = Pt(8)
    p_m2.paragraph_format.space_before = Pt(0)
    r_text = p_m2.add_run(tarea["mentores"])
    set_run_font(r_text, size=11, color=NEGRO)

    # Entregables
    p_e = doc.add_paragraph()
    p_e.paragraph_format.space_after  = Pt(3)
    p_e.paragraph_format.space_before = Pt(0)
    r_e = p_e.add_run("Entregables")
    set_run_font(r_e, size=11, bold=True, color=AZUL)

    for ent in tarea["entregables"]:
        bullet(ent)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────
# PIE
# ─────────────────────────────────────────
separator()
nota(f"ArchiCheck · Plan Startup Chile Ignite · Generado {datetime.date.today().strftime('%d de %B de %Y')}")

# ─────────────────────────────────────────
# GUARDAR
# ─────────────────────────────────────────
output = r"C:\Users\nicolas.estragues\Documents\Claude\archicheck\Startup Chile\Docs 17Ago\ArchiCheck_Plan_Ignite.docx"
doc.save(output)
print(f"Documento guardado: {output}")
